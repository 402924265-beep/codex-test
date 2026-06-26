from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "cost_model_finance_meeting_script_cn_en.docx"


def set_east_asia_font(run_or_style, font_name="Microsoft YaHei"):
    rpr = run_or_style._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_or_style._element.append(rpr)
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def add_pair(doc, cn, en):
    p_cn = doc.add_paragraph(style="Chinese Line")
    r_cn = p_cn.add_run(cn)
    r_cn.font.name = "Calibri"
    set_east_asia_font(r_cn)

    p_en = doc.add_paragraph(style="English Line")
    r_en = p_en.add_run(en)
    r_en.italic = True
    return p_cn, p_en


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    set_east_asia_font(normal)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        set_east_asia_font(style)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if "English Line" not in styles:
        styles.add_style("English Line", WD_STYLE_TYPE.PARAGRAPH)
    style = styles["English Line"]
    style.base_style = normal
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(79, 91, 104)
    style.paragraph_format.space_after = Pt(7)
    style.paragraph_format.left_indent = Inches(0.18)
    style.paragraph_format.line_spacing = 1.10

    if "Chinese Line" not in styles:
        styles.add_style("Chinese Line", WD_STYLE_TYPE.PARAGRAPH)
    style = styles["Chinese Line"]
    style.base_style = normal
    style.font.name = "Calibri"
    set_east_asia_font(style)
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(20, 34, 50)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.10

    if "Note Box" not in styles:
        styles.add_style("Note Box", WD_STYLE_TYPE.PARAGRAPH)
    style = styles["Note Box"]
    style.base_style = normal
    style.font.name = "Calibri"
    set_east_asia_font(style)
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(31, 58, 95)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.left_indent = Inches(0.05)


def build_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("成本分析模型财务团队沟通稿")
    run.font.name = "Calibri"
    set_east_asia_font(run)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Cost Analysis Model Meeting Script | 2026-06-16")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(85, 95, 109)

    p = doc.add_paragraph(style="Note Box")
    p.add_run("会议目标：用简单方式说明模型怎么用，重点讨论第二页预算差异、小科目描述规则、后续功能需求，以及预算是否真的被过程管控。")
    shade_paragraph(p, "F4F6F9")
    p = doc.add_paragraph(style="Note Box")
    p.add_run("Meeting objective: Explain the model in a simple way, focus on page 2 budget variance, account description rules, future functions, and whether budget control really happens during execution.")
    shade_paragraph(p, "F4F6F9")

    sections = [
        ("一、开场 Opening", [
            ("大家下午好，今天我想用很简单的方式，和大家同步一下这个成本分析模型目前怎么用。", "Good afternoon everyone, today I would like to explain in a simple way how this cost analysis model is currently used."),
            ("这个模型不是为了替代财务判断，而是帮助我们更快发现差异、定位科目、记录原因。", "This model is not meant to replace finance judgment; it is meant to help us find variances faster, locate the accounts, and record the reasons."),
            ("我今天重点讲第二页，因为第二页最接近日常预算分析和月度差异解释。", "Today I will mainly focus on the second page, because it is closest to daily budget analysis and monthly variance explanation."),
        ]),
        ("二、模型怎么用 How To Use The Model", [
            ("第一步是导入预测表、实际表和国内财务表，系统会自动生成全年趋势、月度差异和项目因素。", "The first step is to import the forecast file, actual file, and domestic finance file; the system will then generate annual trends, monthly variance, and project factors."),
            ("第一页主要看全年趋势，比如1到5月是实际，6到12月是预测，这样大家可以先看全年方向。", "The first page is mainly for annual trends; for example, January to May are actuals and June to December are forecasts, so we can first see the full-year direction."),
            ("第二页主要看预算月份和差异原因，尤其是导入4+8以后，我们要重点看5月、6月、7月这些预算和实际变化。", "The second page is mainly for budget months and variance reasons; especially after importing 4+8, we should focus on May, June, and July budget and actual changes."),
            ("第三页是降费项目库，它和月度原因分开管理，避免把正式项目和月度解释混在一起。", "The third page is the cost reduction project library, managed separately from monthly reasons so that formal projects and monthly explanations are not mixed together."),
        ]),
        ("三、第二页演示重点 Page 2 Demonstration Focus", [
            ("在第二页，先选择月份，然后看上方的大科目对比，快速判断是哪几个大类在拉高或拉低费用。", "On page 2, first select the month, then look at the category comparison at the top to quickly see which categories are driving cost up or down."),
            ("如果4+8预测没有拆到小科目，系统会自动按大科目汇总比较，避免小科目差异被放大。", "If the 4+8 forecast is not split into detailed accounts, the system compares at category total level to avoid overstating account-level variance."),
            ("再往下看小科目明细，系统会把同比、环比和金额差异放在一起，方便我们直接找原因。", "Then we look at the detailed accounts, where YoY, MoM, and amount variance are shown together so that we can directly identify the reasons."),
            ("对重点差异，我们需要填写原因、责任、行动和预计影响，这样后续复盘时能追踪。", "For key variances, we need to fill in the reason, owner, action, and expected impact, so that we can track them later."),
        ]),
        ("四、请大家重点反馈小科目描述 Account Description Feedback", [
            ("我想请大家重点看一下“科目描述”这一列，现在我先按科目类型做了固定提示。", "I would like everyone to pay special attention to the Account Description column, where I have added fixed prompts by account type."),
            ("比如工作服、餐费、班车这类，建议用单价乘数量来解释。", "For items like work uniforms, meals, and shuttle buses, I suggest explaining them by unit price times quantity."),
            ("比如维修费、备件费、IT费用、运营费这类，建议用清单编号、供应商、金额和用途来解释。", "For repairs, spare parts, IT expenses, and operating expenses, I suggest using list number, supplier, amount, and purpose."),
            ("人工费这块我们先不上传个人明细，只填汇总的数和量，比如人数、工时和产量。", "For labor cost, we will not upload personal details; we will only use summarized figures such as headcount, hours, and production volume."),
            ("这里我想请大家判断一下，这样的描述方式是否够用，哪些科目还需要更准确的模板。", "Here I would like your judgment on whether this description method is enough and which accounts need more accurate templates."),
        ]),
        ("五、请大家提出想要的新功能 Future Function Requests", [
            ("第二个问题是，大家希望这个模型以后还能实现什么功能。", "The second question is what additional functions you would like this model to have in the future."),
            ("比如是否需要附件上传、原因审批、预算预警、周度追踪、自动邮件提醒，或者导出标准汇报模板。", "For example, do we need attachment upload, reason approval, budget alerts, weekly tracking, automatic email reminders, or standard report export."),
            ("如果某个功能能明显减少大家手工整理的时间，我希望今天可以先记录下来。", "If a function can clearly reduce manual preparation time, I would like to record it today."),
        ]),
        ("六、预算是否真的被管控 Budget Control Discussion", [
            ("最后我想讨论一个更实际的问题，就是我们现在有没有真正管控费用预算。", "Finally, I would like to discuss a very practical question: whether we are really controlling the expense budget."),
            ("举个简单例子，如果一个科目预算是5欧元，但实际第一周就花了7欧元，系统可以发现差异，但业务上有没有人马上处理。", "For a simple example, if one account has a budget of 5 euros but spends 7 euros in the first week, the system can find the variance, but does someone act on it immediately in the business process."),
            ("我们需要确认的是，发现超预算以后，谁负责看，什么时候看，超过多少要预警，是否需要暂停或重新审批。", "We need to confirm who reviews it after an over-budget signal, when they review it, what threshold triggers an alert, and whether spending should be paused or re-approved."),
            ("如果只有月底解释原因，那只是分析；如果过程中能预警和处理，才是真正的预算管控。", "If we only explain reasons at month-end, that is analysis; if we can alert and act during the process, that is real budget control."),
        ]),
        ("七、希望会议输出 Expected Meeting Output", [
            ("今天我希望会议结束时，我们能确认三件事。", "By the end of today’s meeting, I hope we can confirm three things."),
            ("第一，小科目描述模板是否合理。", "First, whether the account description templates are reasonable."),
            ("第二，下一步最需要增加哪些功能。", "Second, which functions should be added next."),
            ("第三，费用预算是否要从月度分析升级到过程预警和过程管控。", "Third, whether expense budget management should move from monthly analysis to in-process warning and control."),
            ("如果大家认可这个方向，后续我会继续把模型优化成更适合财务和业务一起使用的工具。", "If everyone agrees with this direction, I will continue improving the model into a tool that finance and business teams can use together."),
        ]),
    ]

    for heading, pairs in sections:
        doc.add_heading(heading, level=1)
        for cn, en in pairs:
            add_pair(doc, cn, en)

    doc.add_heading("会中可以直接问的问题 Questions To Ask Directly", level=1)
    questions = [
        ("大家觉得科目描述这一列，按现在的模板够不够用？", "Do you think the current templates in the Account Description column are enough?"),
        ("哪些科目一定需要附件，哪些科目只需要文字说明？", "Which accounts must have attachments, and which only need text explanation?"),
        ("人工费用只看汇总数和量，大家觉得是否符合合规要求？", "For labor cost, is it compliant enough to use only summarized figures and quantities?"),
        ("如果预算一周内已经超了，我们现在有没有机制及时发现和处理？", "If the budget is already exceeded within one week, do we have a mechanism to detect and handle it in time?"),
        ("未来最优先想加的功能是什么？", "What is the highest-priority function you want to add next?"),
    ]
    for cn, en in questions:
        add_pair(doc, cn, en)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Candy DW Cost Model Meeting Script")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 110, 120)

    core = doc.core_properties
    core.title = "成本分析模型财务团队沟通稿"
    core.subject = "Cost analysis model meeting script"
    core.author = "Candy DW Cost"
    core.keywords = "cost model, finance meeting, budget control"

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
